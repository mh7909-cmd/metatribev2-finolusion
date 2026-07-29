import sqlite3
import csv
import os
import shutil
from datetime import datetime
from pathlib import Path
import config

def init_db():
    """
    Initializes the SQLite database schema if it doesn't exist.
    """
    conn = sqlite3.connect(config.DB_PATH)
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS interactions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            timestamp TEXT NOT NULL,
            question TEXT NOT NULL,
            variant_1 TEXT NOT NULL,
            score_1 REAL NOT NULL,
            variant_2 TEXT NOT NULL,
            score_2 REAL NOT NULL,
            variant_3 TEXT NOT NULL,
            score_3 REAL NOT NULL,
            selected_variant TEXT NOT NULL,
            selected_score REAL NOT NULL,
            audio_path TEXT,
            brain_plot_path TEXT
        )
    """)
    # Database migration: check if brain_plot_path exists, if not, add it
    try:
        cursor.execute("ALTER TABLE interactions ADD COLUMN brain_plot_path TEXT")
    except sqlite3.OperationalError:
        pass  # Column already exists
    conn.commit()
    conn.close()

def log_interaction(question: str, variants_with_scores: list[dict], selected_variant: dict, temp_audio_path: str) -> str:
    """
    Saves a structured log of the interaction:
    1. Copies the temporary audio to a persistent archive location.
    2. Writes to SQLite database.
    3. Writes/Appends to a CSV file.
    Returns the archived audio path.
    """
    init_db()
    
    timestamp_str = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    timestamp_db = datetime.now().isoformat()
    
    # Archive audio file if it exists
    archived_audio_path = ""
    if temp_audio_path and os.path.exists(temp_audio_path):
        dest_filename = f"response_{timestamp_str}.wav"
        dest_path = config.SAVED_AUDIO_DIR / dest_filename
        try:
            shutil.copy(temp_audio_path, dest_path)
            archived_audio_path = str(dest_path)
            print(f"[Logger] Audio archived to {archived_audio_path}")
        except Exception as e:
            print(f"[Logger] Failed to copy audio to archive: {e}")
            archived_audio_path = temp_audio_path
            
    # Unpack variants (expecting exactly 3)
    v1 = variants_with_scores[0]["text"]
    s1 = variants_with_scores[0]["score"]
    v2 = variants_with_scores[1]["text"]
    s2 = variants_with_scores[1]["score"]
    v3 = variants_with_scores[2]["text"]
    s3 = variants_with_scores[2]["score"]
    
    sel_text = selected_variant["text"]
    sel_score = selected_variant["score"]
    brain_plot_path = selected_variant.get("brain_plot_path", "")
    
    # 1. Save to SQLite
    try:
        conn = sqlite3.connect(config.DB_PATH)
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO interactions (
                timestamp, question, 
                variant_1, score_1, 
                variant_2, score_2, 
                variant_3, score_3, 
                selected_variant, selected_score, 
                audio_path, brain_plot_path
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            timestamp_db, question,
            v1, s1,
            v2, s2,
            v3, s3,
            sel_text, sel_score,
            archived_audio_path,
            brain_plot_path
        ))
        conn.commit()
        conn.close()
        print(f"[Logger] SQLite database updated successfully ({config.DB_PATH})")
    except Exception as e:
        print(f"[Logger] SQLite database log failed: {e}")
        
    # 2. Save to CSV
    csv_headers = [
        "timestamp", "question",
        "variant_1", "score_1",
        "variant_2", "score_2",
        "variant_3", "score_3",
        "selected_variant", "selected_score",
        "audio_path", "brain_plot_path"
    ]
    
    write_header = not os.path.exists(config.CSV_PATH)
    
    try:
        with open(config.CSV_PATH, mode="a", newline="", encoding="utf-8") as csv_file:
            writer = csv.writer(csv_file)
            if write_header:
                writer.writerow(csv_headers)
            writer.writerow([
                timestamp_db, question,
                v1, s1,
                v2, s2,
                v3, s3,
                sel_text, sel_score,
                archived_audio_path,
                brain_plot_path
            ])
        print(f"[Logger] CSV log appended successfully ({config.CSV_PATH})")
    except Exception as e:
        print(f"[Logger] CSV log failed: {e}")
        
    # Generate the Word Document report
    report_path = write_docx_report(question, variants_with_scores, selected_variant, archived_audio_path)
        
    return archived_audio_path, report_path

def generate_brain_analysis(phrasing: str, pfc: float, amygdala: float, temporal: float, nacc: float) -> str:
    """
    Calls the NVIDIA LLM to generate a professional, plain-English analysis paragraph
    linking the brain activation metrics (as percentages) to the phrasing style.
    """
    pfc_pct = int(round(pfc * 100))
    amy_pct = int(round(amygdala * 100))
    temp_pct = int(round(temporal * 100))
    nacc_pct = int(round(nacc * 100))

    if not config.NVIDIA_API_KEY:
        return (
            f"This phrasing generated a balanced neural profile, showing a PFC logic score of {pfc_pct}%, "
            f"an Amygdala excitement score of {amy_pct}%, a Temporal cadence score of {temp_pct}%, and a NAcc reward score of {nacc_pct}%."
        )
    try:
        from openai import OpenAI
        client = OpenAI(
            base_url=config.NVIDIA_BASE_URL,
            api_key=config.NVIDIA_API_KEY
        )
        
        prompt = f"""
Analyze the neural activations predicted by the Meta TRIBE v2 model for the following response phrasing.
Phrasing Text: "{phrasing}"

Predicted Region Metrics:
- Prefrontal Cortex (PFC): {pfc_pct}% (Depth, logic, structure)
- Amygdala: {amy_pct}% (Emotion, excitement)
- Temporal (Auditory/Broca): {temp_pct}% (Cadence, rhythm)
- Nucleus Accumbens (NAcc): {nacc_pct}% (Reward, hook)

Write a short, clear, executive analysis paragraph of exactly 2-3 sentences.
Follow these strict instructions:
1. Explain the results in simple, plain-English business terms that are easy to understand. Do NOT use overly complex academic, medical, or "neuroscientific" jargon (avoid words like "cognitive and affective processing", "nuanced interplay", "lateral cortical fMRI", etc.).
2. You MUST explicitly mention the percentages for all 4 regions in the text (e.g. "...resulting in a PFC logic score of {pfc_pct}% and a Temporal score of {temp_pct}%...").
3. Keep it extremely direct: explain how the words/style of the phrasing caused these scores.
4. Do not prefix with "Analysis:" or "Paragraph:". Do not use markdown.
"""
        response = client.chat.completions.create(
            model=config.LLM_MODEL,
            messages=[
                {"role": "system", "content": "You are a professional business advisor interpreting fMRI neural twin simulation results for language phrasings."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=200
        )
        analysis_text = response.choices[0].message.content.strip()
        return analysis_text
    except Exception as e:
        print(f"[Logger] Failed to generate LLM brain analysis: {e}")
        return (
            f"This phrasing generated a balanced neural profile, showing a PFC logic score of {pfc_pct}%, "
            f"an Amygdala excitement score of {amy_pct}%, a Temporal cadence score of {temp_pct}%, and a NAcc reward score of {nacc_pct}%."
        )

def write_docx_report(question: str, variants: list[dict], selected: dict, archived_audio_path: str) -> str:
    """
    Generates a beautifully styled MS Word (.docx) report detailing the run.
    Embeds the leaderboard table and all three unique brain simulation plots side-by-side.
    """
    try:
        import docx
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        reports_dir = Path(config.WORKSPACE_DIR) / "reports"
        reports_dir.mkdir(parents=True, exist_ok=True)
        
        timestamp_str = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        report_file = reports_dir / f"FinoLusion_Cognitive_CFO_Report_{timestamp_str}.docx"
        
        doc = docx.Document()
        
        # Set executive page margins (0.6 inches) for wider data layout
        for section in doc.sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
        
        # --- TITLE BLOCK ---
        title_p = doc.add_paragraph()
        title_run = title_p.add_run("FINOLUSION COGNITIVE CFO REPORT")
        title_run.bold = True
        title_run.font.size = Pt(22)
        title_run.font.color.rgb = RGBColor(11, 25, 44) # Premium dark navy
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        sub_p = doc.add_paragraph()
        sub_run = sub_p.add_run("Neuro-Linguistic Engagement & fMRI Cortical Simulation")
        sub_run.font.size = Pt(13)
        sub_run.italic = True
        sub_run.font.color.rgb = RGBColor(102, 102, 102)
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        meta_p = doc.add_paragraph()
        meta_run = meta_p.add_run(
            f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            f"Subject: Human Neural Twin (TRIBE v2 Backbone)\n"
            f"Audio Response: {os.path.basename(archived_audio_path) if archived_audio_path else 'None'}"
        )
        meta_run.font.size = Pt(10.5)
        meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # --- SECTION 1: QUESTION ---
        doc.add_heading("1. Input Prompt", level=1)
        q_p = doc.add_paragraph()
        q_run = q_p.add_run(f'"{question}"')
        q_run.italic = True
        q_run.font.size = Pt(13)
        q_run.font.color.rgb = RGBColor(102, 102, 102) # Dark grey
        
        # --- SECTION 2: LEADERBOARD ---
        doc.add_heading("2. Response Phrasing Leaderboard", level=1)
        doc.add_paragraph("The table below ranks the three response phrasings by their predicted cognitive engagement:")
        
        # Calculate baseline (lowest score in the batch)
        min_score = min(v["score"] for v in variants)
        
        # Width configurations (Total = 7.3 Inches of printable width)
        col_widths = [Inches(0.6), Inches(3.5), Inches(0.7), Inches(0.7), Inches(0.45), Inches(0.45), Inches(0.45), Inches(0.45)]
        
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Light Shading Accent 1'
        table.allow_autofit = False
        table.autofit = False
        
        # Helper to style cells beautifully and prevent text squishing
        def style_cell(cell, text, width, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, pad=False):
            cell.text = text
            cell.width = width
            p = cell.paragraphs[0]
            p.alignment = align
            # High-end line spacing and cell margin padding for breathing room
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            if pad:
                p.paragraph_format.left_indent = Pt(6)
                p.paragraph_format.right_indent = Pt(6)
            else:
                p.paragraph_format.left_indent = Pt(0)
                p.paragraph_format.right_indent = Pt(0)
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9.5)
                run.bold = bold
                
        hdr_cells = table.rows[0].cells
        headers = ['Rank', 'Phrasing Variant', 'Score', 'Delta', 'PFC', 'Amy', 'Temp', 'NAcc']
        for idx, h_text in enumerate(headers):
            style_cell(hdr_cells[idx], h_text, col_widths[idx], bold=True, align=WD_ALIGN_PARAGRAPH.LEFT if idx == 1 else WD_ALIGN_PARAGRAPH.CENTER, pad=(idx == 1))
            
        # Sort variants by score descending
        sorted_vars = sorted(variants, key=lambda x: x["score"], reverse=True)
        for i, var in enumerate(sorted_vars, 1):
            row_cells = table.add_row().cells
            rank_text = "★ #1\n(Winner)" if i == 1 else f"#{i}"
            style_cell(row_cells[0], rank_text, col_widths[0], bold=(i == 1), align=WD_ALIGN_PARAGRAPH.CENTER, pad=False)
            # Use full text without ellipsis, but justify-aligned for clean boundaries
            style_cell(row_cells[1], var["text"], col_widths[1], align=WD_ALIGN_PARAGRAPH.JUSTIFY, pad=True)
            style_cell(row_cells[2], f"{var['score']:.4f}", col_widths[2], align=WD_ALIGN_PARAGRAPH.CENTER, pad=False)
            
            # Calculate relative delta
            score = var["score"]
            if min_score > 0 and score != min_score:
                rel_delta = ((score - min_score) / min_score) * 100
                rel_delta_str = f"+{rel_delta:.2f}%"
            else:
                rel_delta_str = "Baseline"
            style_cell(row_cells[3], rel_delta_str, col_widths[3], align=WD_ALIGN_PARAGRAPH.CENTER, pad=False)
            
            acts = var.get("activations", {})
            style_cell(row_cells[4], f"{acts.get('PFC', 0.0):.2f}", col_widths[4], align=WD_ALIGN_PARAGRAPH.CENTER, pad=False)
            style_cell(row_cells[5], f"{acts.get('Amygdala', 0.0):.2f}", col_widths[5], align=WD_ALIGN_PARAGRAPH.CENTER, pad=False)
            style_cell(row_cells[6], f"{acts.get('Temporal', 0.0):.2f}", col_widths[6], align=WD_ALIGN_PARAGRAPH.CENTER, pad=False)
            style_cell(row_cells[7], f"{acts.get('NAcc', 0.0):.2f}", col_widths[7], align=WD_ALIGN_PARAGRAPH.CENTER, pad=False)

        # Force lock column widths across all rows and cells
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = width
                
        doc.add_paragraph() # Spacing
        
        # --- SECTION 3: BRAIN VISUALIZATIONS ---
        doc.add_heading("3. 3D Cortical fMRI Simulation Maps", level=1)
        doc.add_paragraph(
            "Below are the 3D lateral projections of predicted BOLD signals across the cortical surface of the brain "
            "for each phrasing variant. Warm colors (red/yellow) represent high activation, and cool colors (blue/green) represent low activation:"
        )
        
        for i, var in enumerate(sorted_vars, 1):
            rank_str = "★ #1 (Winner)" if i == 1 else f"#{i}"
            doc.add_heading(f"Variant {rank_str} (Score: {var['score']:.4f})", level=2)
            
            p_desc = doc.add_paragraph()
            p_desc.add_run(f"Phrasing: ").bold = True
            p_desc.add_run(f"\"{var['text']}\"")
            
            plot_path = var.get("brain_plot_path")
            if plot_path and os.path.exists(plot_path):
                # Add picture centered
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = p_img.add_run()
                run_img.add_picture(plot_path, width=Inches(5.0))
            else:
                p_err = doc.add_paragraph()
                p_err.add_run("No 3D brain plot available for this variant.").italic = True
                
            # Dynamic brain analysis paragraph below the image
            acts = var.get("activations", {})
            pfc_val = acts.get("PFC", 0.0)
            amy_val = acts.get("Amygdala", 0.0)
            temp_val = acts.get("Temporal", 0.0)
            nacc_val = acts.get("NAcc", 0.0)
            
            print(f"[Logger] Generating dynamic brain analysis for Variant {i}...")
            analysis_text = generate_brain_analysis(var["text"], pfc_val, amy_val, temp_val, nacc_val)
            
            p_analysis = doc.add_paragraph()
            p_analysis.paragraph_format.space_before = Pt(6)
            p_analysis.paragraph_format.space_after = Pt(12)
            run_lbl = p_analysis.add_run("Neural Activation Analysis: ")
            run_lbl.bold = True
            p_analysis.add_run(analysis_text)
            
        doc.save(report_file)
        print(f"[Logger] MS Word report created successfully: {report_file}")
        return str(report_file)
    except Exception as e:
        print(f"[Logger] Failed to write DOCX report: {e}")
        return ""

if __name__ == "__main__":
    # Test logger
    init_db()
    test_q = "What is the capital of France?"
    test_v = [
        {"text": "V1 description", "score": 0.5},
        {"text": "V2 description", "score": 0.6},
        {"text": "V3 description", "score": 0.7}
    ]
    test_sel = {"text": "V3 description", "score": 0.7}
    log_interaction(test_q, test_v, test_sel, "")
